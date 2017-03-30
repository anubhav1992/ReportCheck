import docx
import os
import re
import win32com.client.gencache
import sys


#if sys.argv[1] == "-h":
#    print '''\nReportCheck\n---------------------\nPlease give the full path of your "docm." report along with the report name.
#    \nExample :\n------------\npython ReportCheck.py\nFull_Path_To_Report\Report_Name.docm'''
#    sys.exit(0)

try :
    reportPath = raw_input("Please Enter The Report Path : \n")
    reportPath = os.path.normpath(reportPath)
    doc = docx.Document(reportPath)
    #print len(doc.paragraphs)
    l = []
    for i in range(0, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            l.append(doc.paragraphs[i].text)
        else:
            pass

    clientName = l[0]
    appName = l[1]
    testType = l[2]

    if l[7].startswith("For more information"):
        asrName = l[8]
        asrRole = l[9]
        toName = l[10]
        toRole = l[11]
        mgrName = l[12]
        mgrRole = l[13]
    else:
        asrName = l[7]
        asrRole = l[8]
        toName = l[9]
        toRole = l[10]
        mgrName = l[11]
        mgrRole = l[12]
    print "************************************************************************************************************"
    print "\nScan Details\n=============="
    print "Client Name : " + clientName + "\n"
    print "Application Name : " + appName + "\n"
    print "Test Type : " + testType + "\n"
    if l[6].startswith("For more information"):
        print "Client Address : None"+"\n"
    else:
        clientAddress = l[6]
        print "Client Address : "+ clientAddress+"\n"
    print "************************************************************************************************************"
    print "Assessor Details\n=================="
    print "Assessor's Name : " + asrName
    print "Assessor's Designation : "+ asrRole + "\n"
    print "TO Name : "+ toName
    print "TO Designation : "+ toRole + "\n"
    print "Manager Name : "+ mgrName
    print "Manager Designation : "+ mgrRole + "\n"
    print "************************************************************************************************************"
    firstLine = 0
    lastLine = 0
    sumFindings = 0
    exError = 0
    intError = 0
    totalFind = 0
    for i in range(0,len(l)):
        if l[i].startswith("Executive Summary"):
            exSum = l[i+1]
            count = sum(1 for _ in re.finditer(r'\b%s\b' % re.escape(appName), exSum))
            if not exSum.startswith(clientName):
                print "\n\033[1;91mError : Client Name mismatched in executive summary. Please correct it.\033[0;m"
                exError += 1
            if (not appName in l[i + 1]) or (not count == 2):
                print "\n\033[1;91mError : Application Name mismatched in executive summary. Please correct it.\033[0;m"
                exError += 1
            if not testType in l[i + 1]:
                print "\n\033[1;91mError : Offering Name mismatched in executive summary. Please correct it.\033[0;m"
                exError += 1
            if exError == 0:
                print "\n\033[1;92mExecutive Summary is correct. :)\033[0;m\n"
        else:
            pass
        if "findings characterized as follows:" in l[i]:
            totalFind += 1
            firstLine =  i+1
            p = l[i].split()
            for k in range(0,len(p)):
                if p[k].startswith("identified"):
                    sumFindings = p[k+1]
                else:
                    pass
        else:
            pass

        if ("While performing the assessment" in l[i]) and (totalFind != 0):
            lastLine =  i-1
        else:
            pass
        if l[i].startswith("Introduction"):
            intro = l[i + 1]
            count = sum(1 for _ in re.finditer(r'\b%s\b' % re.escape(appName), intro))
            if not intro.startswith(clientName):
                print "\n\033[1;91mError : Client Name mismatched in introduction part. Please correct it.\033[0;m"
                intError += 1
            if (not appName in l[i + 1]) or (not count == 2):
                print "\n\033[1;91mError : Application Name mismatched in introduction part. Please correct it.\033[0;m"
                intError += 1
            if not testType in l[i + 1]:
                intError += 1
                print "\n\033[1;91mError : Offering Name mismatched in introduction part. Please correct it.\033[0;m"
            if intError == 0:
                print "\n\033[1;92mIntroduction is correct. :)\033[0;m\n"
        else:
            pass
    print "************************************************************************************************************"
    if totalFind != 0:
        print "Findings Categories : \n====================="
        a =[]
        for line in range(firstLine,lastLine+1):
            print l[line]
            a.append(l[line].split())

        Sum = 0

        for t in range(0,len(a)):
            Sum += int(a[t][0])

        print "\nSum of findings is : ",format(Sum)
        print "\nTotal findings written in finding categories : "+ sumFindings
        print "************************************************************************************************************"
        if Sum == int(sumFindings):
            print "\n\033[1;92mNumber of findings are matched with the Total of findings.\033[0;m\n"
            print "************************************************************************************************************"
        else:
            print "\n\033[1;91mThere is a mismatch between your Total Findings and Finding Categories.\nPlease correct it.\033[0;m"
    else :
        print "No finding categories are found or may be you've wrong written your Executive Summary. Please Check."
    print "************************************************************************************************************"

#Checking for the Grammatical Errors and Spelling Errors

    check = raw_input("\nDo you want to check the Grammar and Spelling Errors ?\nPlease press Y for yes or Enter for skip.\n")
    check = check.upper()
    if check == "Y":
        print "************************************************************************************************************"
        print "Checking the Grammar Errors.......\n"

        wdDoNotSaveChanges = 0

        app = win32com.client.gencache.EnsureDispatch('Word.Application')
        doc = app.Documents.Open(reportPath)

        grmError = []
        splError = []

        if doc.GrammaticalErrors.Count > 0:
            for err in doc.GrammaticalErrors:
                grmError.append(err.Text)
        if doc.SpellingErrors.Count > 0:
            for err in doc.SpellingErrors:
                splError.append(err.Text)

        if not grmError is None:
            choice = raw_input('\033[1;41mYou have %d grammatical errors in the report, do you want to see it ?\033[0;m\nPress Y for yes or Enter for skip.\n'% (
                doc.GrammaticalErrors.Count))
            choice = choice.upper()
            if choice == "Y":
                print "************************************************************************************************************"
                print "Printing Grammatical Errors........\n"
                for i in range(0, len(grmError)):
                    print "Grammatical Error in : " + grmError[i].encode('utf-8')
            else:
                pass
        print "************************************************************************************************************"
        print "Checking for Spelling Errors........\n"
        if not splError is None:
            print "************************************************************************************************************"
            choice = raw_input(
                "\033[1;41mYou have %d spelling errors in the report, do you want to see it ?\033[0;m\nPress Y for yes or Enter for skip.\n" % (
                doc.SpellingErrors.Count))
            choice = choice.upper()
            if choice == "Y":
                print "************************************************************************************************************"
                print "Printing Spelling Errors........\n"
                for i in range(0, len(splError)):
                    print "Spelling Error in : " + splError[i].encode('utf-8')
                print "************************************************************************************************************"
            else:
                pass

        app.Quit(wdDoNotSaveChanges)
    else:
        print "************************************************************************************************************"
        print "\nYou skipped the Grammatical and Spelling checking.\n"
        print "************************************************************************************************************"
except:
    print "\nSorry ! Some Error Occur.\nMay be the report is not present at the path or you entered the wrong format.\nYou can use -h option for help.\n(python ReportCheck.py -h)"


"""
---------------------------
***************************
Anubhav Sharma
Mail: ansharma@cigital.com
***************************
---------------------------
"""
